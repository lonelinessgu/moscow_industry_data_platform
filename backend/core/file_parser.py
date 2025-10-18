import pandas as pd
from docx import Document
import pdfplumber
from tortoise.exceptions import IntegrityError
from backend.models.companies import Company
import logging

logger = logging.getLogger(__name__)
def normalize_header(header):
    """Преобразует заголовок столбца в формат, соответствующий полям модели."""
    return header.lower().replace(" ", "_").replace("-", "_").replace("(", "").replace(")", "").replace(",", "").replace(".", "")

# Соответствие между заголовками столбцов в файле и полями модели
COLUMN_MAPPING = {
    'инн': 'inn',
    'наименование_организации': 'name',
    'полное_наименование_организации': 'full_name',
    'статус': 'status',
    'юридический_адрес': 'legal_address',
    'адрес_производства': 'production_address',
    'адрес_дополнительной_площадки': 'additional_site_address',
    'основная_отрасль': 'main_industry',
    'подотрасль_основная': 'sub_industry',
    'основной_оквэд': 'main_okved',
    'вид_деятельности_по_основному_оквэд': 'main_okved_activity',
    'производственный_оквэд': 'production_okved',
    'дата_регистрации': 'registration_date',
    'руководитель': 'director',
    'головная_организация': 'head_company',
    'инн_головной_организации': 'head_company_inn',
    'контактные_данные_руководства': 'director_contacts',
    'контакт_сотрудника_организации': 'employee_contacts',
    'контактные_данные_ответственного_по_чс': 'cs_responsible_contacts',
    'сайт': 'website',
    'электронная_почта': 'email',
    'данные_об_оказанных_мерах_поддержки': 'support_measures',
    'наличие_особого_статуса': 'special_status',
    'статус_мсп': 'msp_status',
    'выручка_предприятия_тыс_руб_2022': 'revenue_2022',
    'выручка_предприятия_тыс_руб_2023': 'revenue_2023',
    'выручка_предприятия_тыс_руб_2024': 'revenue_2024',
    'чистая_прибыль_убыток_тыс_руб_2022': 'net_profit_2022',
    'чистая_прибыль_убыток_тыс_руб_2023': 'net_profit_2023',
    'чистая_прибыль_убыток_тыс_руб_2024': 'net_profit_2024',
    'среднесписочная_численность_персонала_всего_по_компании_чел_2022': 'staff_count_total_2022',
    'среднесписочная_численность_персонала_всего_по_компании_чел_2023': 'staff_count_total_2023',
    'среднесписочная_численность_персонала_всего_по_компании_чел_2024': 'staff_count_total_2024',
    'среднесписочная_численность_персонала_работающего_в_москве_чел_2022': 'staff_count_moscow_2022',
    'среднесписочная_численность_персонала_работающего_в_москве_чел_2023': 'staff_count_moscow_2023',
    'среднесписочная_численность_персонала_работающего_в_москве_чел_2024': 'staff_count_moscow_2024',
    'фонд_оплаты_труда_всех_сотрудников_организации_тыс_руб_2022': 'payroll_total_2022',
    'фонд_оплаты_труда_всех_сотрудников_организации_тыс_руб_2023': 'payroll_total_2023',
    'фонд_оплаты_труда_всех_сотрудников_организации_тыс_руб_2024': 'payroll_total_2024',
    'фонд_оплаты_труда_сотрудников_работающих_в_москве_тыс_руб_2022': 'payroll_moscow_2022',
    'фонд_оплаты_труда_сотрудников_работающих_в_москве_тыс_руб_2023': 'payroll_moscow_2023',
    'фонд_оплаты_труда_сотрудников_работающих_в_москве_тыс_руб_2024': 'payroll_moscow_2024',
    'средняя_з_п_всех_сотрудников_организации_тыс_руб_2022': 'avg_salary_total_2022',
    'средняя_з_п_всех_сотрудников_организации_тыс_руб_2023': 'avg_salary_total_2023',
    'средняя_з_п_всех_сотрудников_организации_тыс_руб_2024': 'avg_salary_total_2024',
    'средняя_з_п_сотрудников_работающих_в_москве_тыс_руб_2022': 'avg_salary_moscow_2022',
    'средняя_з_п_сотрудников_работающих_в_москве_тыс_руб_2023': 'avg_salary_moscow_2023',
    'средняя_з_п_сотрудников_работающих_в_москве_тыс_руб_2024': 'avg_salary_moscow_2024',
    'налоги_уплаченные_в_бюджет_москвы_без_акцизов_тыс_руб_2022': 'taxes_paid_moscow_2022',
    'налоги_уплаченные_в_бюджет_москвы_без_акцизов_тыс_руб_2023': 'taxes_paid_moscow_2023',
    'налоги_уплаченные_в_бюджет_москвы_без_акцизов_тыс_руб_2024': 'taxes_paid_moscow_2024',
    'налог_на_прибыль_тыс_руб_2022': 'profit_tax_2022',
    'налог_на_прибыль_тыс_руб_2023': 'profit_tax_2023',
    'налог_на_прибыль_тыс_руб_2024': 'profit_tax_2024',
    'налог_на_имущество_тыс_руб_2022': 'property_tax_2022',
    'налог_на_имущество_тыс_руб_2023': 'property_tax_2023',
    'налог_на_имущество_тыс_руб_2024': 'property_tax_2024',
    'налог_на_землю_тыс_руб_2022': 'land_tax_2022',
    'налог_на_землю_тыс_руб_2023': 'land_tax_2023',
    'налог_на_землю_тыс_руб_2024': 'land_tax_2024',
    'ндфл_тыс_руб_2022': 'ndfl_2022',
    'ндфл_тыс_руб_2023': 'ndfl_2023',
    'ндфл_тыс_руб_2024': 'ndfl_2024',
    'транспортный_налог_тыс_руб_2022': 'transport_tax_2022',
    'транспортный_налог_тыс_руб_2023': 'transport_tax_2023',
    'транспортный_налог_тыс_руб_2024': 'transport_tax_2024',
    'прочие_налоги_2022': 'other_taxes_2022',
    'прочие_налоги_2023': 'other_taxes_2023',
    'прочие_налоги_2024': 'other_taxes_2024',
    'акцизы_тыс_руб_2022': 'excise_2022',
    'акцизы_тыс_руб_2023': 'excise_2023',
    'акцизы_тыс_руб_2024': 'excise_2024',
    'инвестиции_в_мск_2022_тыс_руб': 'investments_moscow_2022',
    'инвестиции_в_мск_2023_тыс_руб': 'investments_moscow_2023',
    'инвестиции_в_мск_2024_тыс_руб': 'investments_moscow_2024',
    'объем_экспорта_тыс_руб_2022': 'export_volume_2022',
    'объем_экспорта_тыс_руб_2023': 'export_volume_2023',
    'объем_экспорта_тыс_руб_2024': 'export_volume_2024',
    'кадастровый_номер_зу': 'cadastral_number_land',
    'площадь_зу': 'area_land',
    'вид_разрешенного_использования_зу': 'permitted_use_land',
    'вид_собственности_зу': 'ownership_type_land',
    'собственник_зу': 'land_owner',
    'кадастровый_номер_окса': 'cadastral_number_oks',
    'площадь_оксов': 'area_oks',
    'вид_разрешенного_использования_оксов': 'permitted_use_oks',
    'тип_строения_и_цель_использования': 'building_type',
    'вид_собственности_оксов': 'ownership_type_oks',
    'собственник_оксов': 'oks_owner',
    'площадь_производственных_помещений_кв_м': 'production_area_sqm',
    'стандартизированная_продукция': 'standardized_product',
    'название_виды_производимой_продукции': 'product_types',
    'перечень_производимой_продукции_по_кодам_окпд_2': 'okpd2_products',
    'перечень_производимой_продукции_по_типам_и_сегментам': 'product_segments',
    'каталог_продукции': 'product_catalog',
    'наличие_госзаказа': 'state_contract',
    'уровень_загрузки_производственных_мощностей': 'capacity_utilization',
    'наличие_поставок_продукции_на_экспорт': 'export_supply',
    'объем_экспорта_млн_руб_за_предыдущий_календарный_год': 'export_volume_prev_year',
    'перечень_государств_импортеров': 'export_countries',
    'координаты_юридического_адреса': 'legal_address_coords',
    'координаты_адреса_производства': 'production_address_coords',
    'координаты_адреса_дополнительной_площадки': 'additional_site_coords',
    'координаты_широта': 'latitude',
    'координаты_долгота': 'longitude',
    'округ': 'district',
    'район': 'area',
}

def extract_data_from_df(df):
    """Извлекает данные из DataFrame, учитывая заголовки столбцов."""
    if df.empty:
        logger.info("DataFrame пуст.")
        return []

    # Предполагаем, что первая строка - заголовки
    headers_raw = df.columns.tolist() if df.columns.name is None else df.iloc[0].tolist()

    # Нормализуем заголовки
    normalized_headers = [normalize_header(str(h)) for h in headers_raw]

    logger.info(f"Обнаруженные заголовки: {normalized_headers}")
    logger.info(f"Количество столбцов: {len(normalized_headers)}")

    records = []
    # Проходим по строкам данных (начиная с индекса 1, если первая строка была заголовком)
    # Или с индекса 0, если pandas правильно прочитал заголовки
    for idx, row in df.iterrows():
        record = {}
        for col_idx, value in enumerate(row):
            if pd.isna(value):
                value = None
            else:
                value = str(value).strip()
                if value == "":
                    value = None

            if col_idx < len(normalized_headers):
                header = normalized_headers[col_idx]
                # Проверяем, есть ли заголовок в нашем сопоставлении
                model_field = COLUMN_MAPPING.get(header)
                if model_field:
                    record[model_field] = value
        records.append(record)
    return records


async def process_excel(df):
    """Обрабатывает DataFrame и сохраняет данные в базу."""
    logger.info("Начало обработки Excel файла.")
    records = extract_data_from_df(df)
    total_processed = len(records)
    created_count = 0
    skipped_count = 0

    logger.info(f"Извлечено {total_processed} строк данных для обработки.")

    for i, record in enumerate(records):
        if not record: # Пропускаем пустые строки
            logger.debug(f"Строка {i} пуста, пропускается.")
            skipped_count += 1
            continue

        logger.debug(f"Обработка строки {i}: {record}")
        try:
            # Создаем экземпляр модели, передавая только сопоставленные поля
            company_instance = Company(**record)
            await company_instance.save()
            created_count += 1
            logger.debug(f"Запись для {record.get('name', 'N/A')} успешно сохранена.")
        except IntegrityError as e:
            logger.warning(f"Ошибка IntegrityError при сохранении строки {i}: {str(e)}. Запись пропущена.")
            skipped_count += 1
        except Exception as e:
            logger.error(f"Неожиданная ошибка при сохранении строки {i}: {str(e)}", exc_info=True)
            skipped_count += 1

    logger.info(f"Обработка Excel файла завершена. Обработано: {total_processed}, создано: {created_count}, пропущено: {skipped_count}.")
    return f"Обработано {total_processed} строк, создано {created_count} записей в базе."

def process_word(path):
    doc = Document(path)
    result_parts = []
    table_num = 0
    for table in doc.tables:
        table_num += 1
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            result_parts.append(" | ".join(cells))
    if table_num == 0:
        for para in doc.paragraphs:
            if para.text.strip():
                result_parts.append(para.text)
    return "\n".join(result_parts)

def process_pdf(path):
    result_parts = []
    found_tables = False
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            if tables:
                found_tables = True
                for table in tables:
                    for row in table:
                        clean = [str(c).strip() if c is not None else "" for c in row]
                        result_parts.append(" | ".join(clean))
    if not found_tables:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text and text.strip():
                    result_parts.append(text)
    return "\n".join(result_parts)